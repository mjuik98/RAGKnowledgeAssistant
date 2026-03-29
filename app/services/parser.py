from __future__ import annotations

import hashlib
from pathlib import Path

from app.models.schemas import ParsedDocument, ParsedPage


def compute_checksum(file_path: Path) -> str:
    hasher = hashlib.sha256()
    with file_path.open("rb") as f:
        for block in iter(lambda: f.read(8192), b""):
            hasher.update(block)
    return hasher.hexdigest()


class DocumentParser:
    supported_suffixes = {".pdf", ".txt", ".docx"}

    def parse(self, file_path: Path) -> ParsedDocument:
        suffix = file_path.suffix.lower()
        if suffix not in self.supported_suffixes:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {suffix}")

        if suffix == ".pdf":
            pages = self._parse_pdf(file_path)
        elif suffix == ".txt":
            pages = self._parse_txt(file_path)
        elif suffix == ".docx":
            pages = self._parse_docx(file_path)
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {suffix}")

        return ParsedDocument(
            title=file_path.stem,
            original_filename=file_path.name,
            source_path=str(file_path),
            source_type=suffix.replace(".", ""),
            checksum=compute_checksum(file_path),
            pages=pages,
        )

    def _parse_pdf(self, file_path: Path) -> list[ParsedPage]:
        try:
            import fitz
        except ImportError as e:
            raise RuntimeError("PDF 파싱을 위해 pymupdf 설치가 필요합니다.") from e

        pages: list[ParsedPage] = []
        with fitz.open(file_path) as doc:
            for idx, page in enumerate(doc):
                text = page.get_text("text").strip()
                pages.append(ParsedPage(page_number=idx + 1, text=text))
        return pages

    def _parse_txt(self, file_path: Path) -> list[ParsedPage]:
        text = file_path.read_text(encoding="utf-8")
        return [ParsedPage(page_number=1, text=text.strip())]

    def _parse_docx(self, file_path: Path) -> list[ParsedPage]:
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise RuntimeError("DOCX 파싱을 위해 python-docx 설치가 필요합니다.") from e

        document = DocxDocument(file_path)
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        return [ParsedPage(page_number=1, text=text.strip())]
